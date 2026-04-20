"""
Gera o relatório do Módulo 3 em formato .docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Margens ──────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(3)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2)

# ── Fonte padrão ─────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font  = style.font
font.name = 'Arial'
font.size = Pt(12)

def add_heading(text, level=1, bold=True, size=14, color=None, center=False, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    run.font.name = 'Arial'
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_body(text, justify=True, size=12, space_after=6, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.bold   = bold
    run.italic = italic
    return p

def add_code(lines):
    """Bloco de código em fonte monoespaçada."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent  = Cm(1)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Cm(level * 0.5 + 0.5)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    return p

def add_table_row(table, cells_data, bold=False, bg=None):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.bold = bold
        if bg:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  bg)
            tcPr.append(shd)
    return row

def hline():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pb.append(bottom)
    pPr.append(pb)

# ══════════════════════════════════════════════════════════════════════════════
# CABEÇALHO
# ══════════════════════════════════════════════════════════════════════════════
add_heading('AVALIAÇÃO DO MÓDULO 3', level=1, size=14, center=True, space_before=0)
add_heading('Banco de Dados e Controle de Versão', level=1, size=13, center=True,
            bold=False, space_before=2)
hline()

meta = [
    ('Projeto de Extensão',    'Programa de Extensão UFMS Digital (95DX7.200525)'),
    ('Nome completo',          'Nilton Dobes Bakargi'),
    ('Disciplina',             'Checkout de Presença e Avaliação da Aprendizagem do Módulo 3\n— Relatório Banco de Dados e Controle de Versão'),
    ('Semestre letivo',        '8º Semestre'),
    ('Professor Especialista', 'Awdren de Lima Fontão'),
    ('Tutor',                  'Elton Rivanor Sena dos Santos'),
    ('Público-alvo',           'Colaboradores de campo, especialistas botânicos e gestores acadêmicos da UFMS e UEMS'),
    ('Local de realização',    'Universidade Federal de Mato Grosso do Sul (UFMS) — Campo Grande, MS'),
]
for label, value in meta:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    r1 = p.add_run(f'{label}: ')
    r1.bold      = True
    r1.font.name = 'Arial'
    r1.font.size = Pt(11)
    r2 = p.add_run(value)
    r2.font.name = 'Arial'
    r2.font.size = Pt(11)

hline()
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# TÍTULO
# ══════════════════════════════════════════════════════════════════════════════
add_heading('PENOMATO: MODELAGEM E VERSIONAMENTO DE UM BANCO DE DADOS\nPARA CATALOGAÇÃO MORFOLÓGICA DE ESPÉCIES DO CERRADO',
            size=13, center=True, space_before=8, space_after=10)

# ══════════════════════════════════════════════════════════════════════════════
# RESUMO
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Resumo', size=12, space_before=8, space_after=4)
add_body(
    'Este relatório apresenta a modelagem, implementação e o gerenciamento por controle de versão '
    'do banco de dados do sistema Penomato — uma plataforma web voltada à catalogação científica '
    'de espécies arbóreas do Cerrado, desenvolvida em parceria entre a UFMS e a UEMS no âmbito '
    'da Engenharia Florestal. O banco de dados, implementado em MariaDB, é composto por nove '
    'tabelas relacionais que cobrem o ciclo completo de uma espécie: do cadastro inicial pelo gestor '
    'até a publicação do artigo científico revisado por especialista. O modelo contempla entidades '
    'para usuários com controle de permissões por categoria, espécies com progressão de status '
    'auditável, características morfológicas organizadas por parte da planta, exemplares físicos '
    'de campo com geolocalização, imagens vinculadas a exemplares aprovados, histórico de auditoria, '
    'sugestões da comunidade e controles de segurança de acesso. O projeto é versionado com Git '
    'desde o commit inicial, com 114 commits documentados e mensagens descritivas seguindo boas '
    'práticas de versionamento semântico.'
)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
r = p.add_run('Palavras-chave: ')
r.bold = True; r.font.name = 'Arial'; r.font.size = Pt(12)
r2 = p.add_run('Banco de Dados Relacional. Controle de Versão. Cerrado. Catalogação Botânica. SQL.')
r2.font.name = 'Arial'; r2.font.size = Pt(12)

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUÇÃO
# ══════════════════════════════════════════════════════════════════════════════
add_heading('1. Introdução', size=13, space_before=12)
add_body(
    'A conservação da biodiversidade do Cerrado — segundo maior bioma do Brasil e um dos '
    '34 hotspots de biodiversidade do planeta — demanda sistemas de informação capazes de organizar, '
    'validar e disseminar dados sobre sua flora de forma científica e colaborativa. Nesse contexto '
    'surge o Penomato, um sistema web desenvolvido para apoiar pesquisadores e alunos da área de '
    'Engenharia Florestal da UEMS (Universidade Estadual de Mato Grosso do Sul) na tarefa de '
    'catalogar morfologicamente espécies arbóreas do Cerrado.'
)
add_body(
    'O desafio central é garantir que os dados inseridos no sistema — atributos morfológicos de '
    'folha, flor, fruto, caule e semente — sejam confiáveis, rastreáveis e validados cientificamente. '
    'Para isso, a base de dados precisa refletir fielmente o fluxo de trabalho científico: desde a '
    'coleta de dados em campo, passando pela revisão de um especialista, até a publicação de um '
    'artigo com créditos a todos os contribuidores.'
)
add_body(
    'A escolha pelo banco de dados relacional (MariaDB) é justificada pela natureza fortemente '
    'estruturada dos dados botânicos e pelas relações complexas entre entidades. O controle de '
    'versão com Git é adotado desde o início do projeto para garantir rastreabilidade de todas '
    'as decisões de desenvolvimento, permitindo reverter alterações, documentar a evolução do '
    'sistema e manter um histórico completo que serve como documentação técnica do projeto.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. OBJETIVO GERAL
# ══════════════════════════════════════════════════════════════════════════════
add_heading('2. Objetivo Geral', size=13, space_before=12)
add_body(
    'Desenvolver e documentar o banco de dados relacional do sistema Penomato, implementando o '
    'esquema de tabelas, restrições de integridade e operações de manipulação de dados, gerenciando '
    'toda a evolução do código por meio de controle de versão com Git.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. OBJETIVOS ESPECÍFICOS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('3. Objetivos Específicos', size=13, space_before=12)
objetivos = [
    'Projetar o modelo de dados do Penomato, definindo entidades, atributos, relacionamentos e '
    'restrições que representem o fluxo científico de catalogação de espécies.',
    'Implementar o banco de dados em MariaDB utilizando SQL, aplicando normalização, chaves '
    'primárias, chaves estrangeiras, índices e tipos de dados adequados a cada domínio.',
    'Executar operações DML (inserção, atualização, remoção e consultas) que reflitam as ações '
    'reais do sistema, como o cadastro de espécies, envio de imagens e aprovação de exemplares.',
    'Gerenciar o código do projeto com Git, utilizando commits frequentes e com mensagens '
    'descritivas, organizando a evolução do banco de dados em arquivos SQL versionados.',
    'Publicar o repositório no GitHub, tornando o histórico de desenvolvimento auditável e o '
    'código acessível para a comunidade acadêmica.',
]
for obj in objetivos:
    add_bullet(obj)

# ══════════════════════════════════════════════════════════════════════════════
# 4. JUSTIFICATIVA
# ══════════════════════════════════════════════════════════════════════════════
add_heading('4. Justificativa e Delimitação do Problema', size=13, space_before=12)
add_body(
    'A documentação científica de espécies do Cerrado é realizada, em grande parte, por meio de '
    'fichas em papel ou planilhas isoladas, sem integração entre coleta de campo, revisão '
    'especializada e publicação. Esse processo fragmentado dificulta a colaboração, compromete a '
    'rastreabilidade dos dados e impede a construção de um acervo público e confiável.'
)
add_body(
    'O Penomato resolve esse problema ao centralizar todo o fluxo em um único sistema com banco '
    'de dados relacional. A escolha do MariaDB é justificada pela familiaridade da equipe, pelo '
    'suporte nativo ao XAMPP (ambiente de desenvolvimento local) e pela adequação ao porte do '
    'projeto. O uso do Git como sistema de controle de versão é essencial para um projeto em '
    'desenvolvimento ativo: permite rastrear cada alteração no esquema do banco de dados, '
    'documentar o motivo de cada mudança nas mensagens de commit e manter histórico completo '
    'que serve como documentação técnica.'
)
add_body(
    'A delimitação do problema concentra-se nas espécies arbóreas do Cerrado registradas pela '
    'UEMS no Mato Grosso do Sul, com foco nos biomas Cerrado, Pantanal e Mata Atlântica do estado.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. FUNDAMENTAÇÃO TEÓRICA
# ══════════════════════════════════════════════════════════════════════════════
add_heading('5. Fundamentação Teórica', size=13, space_before=12)

add_heading('5.1 Modelagem de Banco de Dados Relacional', size=12, space_before=8, space_after=4)
add_body(
    'O modelo relacional, proposto por Edgar F. Codd (1970), organiza dados em tabelas compostas '
    'por linhas (tuplas) e colunas (atributos). O modelo Entidade-Relacionamento (ER), desenvolvido '
    'por Chen (1976), é a ferramenta padrão para projetar a estrutura conceitual antes da '
    'implementação física. No Penomato foram identificadas as entidades: Usuário, Espécie, '
    'Características Morfológicas, Exemplar, Imagem de Parte, Histórico de Alterações e Sugestão.'
)

add_heading('5.2 Normalização', size=12, space_before=8, space_after=4)
add_body(
    'A normalização reduz redundâncias e anomalias de inserção, atualização e exclusão (DATE, 2003). '
    'O banco do Penomato aplica as três primeiras formas normais: 1FN (atributos atômicos), '
    '2FN (dependência total da chave primária) e 3FN (sem dependências transitivas).'
)

add_heading('5.3 SQL — Structured Query Language', size=12, space_before=8, space_after=4)
add_body(
    'SQL é a linguagem padrão para criação e manipulação de bancos de dados relacionais, '
    'padronizada pela ISO/IEC 9075. É dividida em DDL (CREATE, ALTER, DROP), '
    'DML (INSERT, UPDATE, DELETE, SELECT) e DCL (GRANT, REVOKE).'
)

add_heading('5.4 Controle de Versão com Git', size=12, space_before=8, space_after=4)
add_body(
    'Git é um sistema de controle de versão distribuído criado por Linus Torvalds (2005). '
    'Permite rastrear cada alteração no código, reverter mudanças indesejadas e colaborar de '
    'forma organizada (CHACON; STRAUB, 2014). Boas práticas incluem commits atômicos com '
    'mensagens no padrão Conventional Commits e uso de branches para funcionalidades isoladas.'
)

add_heading('Referências da Seção', size=11, bold=False, space_before=6, space_after=4)
refs = [
    'CHEN, P. P. The entity-relationship model. ACM Transactions on Database Systems, v. 1, n. 1, p. 9–36, 1976.',
    'CHACON, S.; STRAUB, B. Pro Git. 2. ed. New York: Apress, 2014.',
    'CODD, E. F. A relational model of data for large shared data banks. Communications of the ACM, v. 13, n. 6, p. 377–387, 1970.',
    'DATE, C. J. Introdução a sistemas de banco de dados. 8. ed. Rio de Janeiro: Campus, 2003.',
    'ELMASRI, R.; NAVATHE, S. B. Sistemas de banco de dados. 7. ed. São Paulo: Pearson, 2018.',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    run = p.add_run(ref)
    run.font.name = 'Arial'
    run.font.size = Pt(10)

# ══════════════════════════════════════════════════════════════════════════════
# 6. METODOLOGIA
# ══════════════════════════════════════════════════════════════════════════════
add_heading('6. Metodologia', size=13, space_before=14)

# ── 6.1 Modelagem ─────────────────────────────────────────────────────────────
add_heading('6.1 Modelagem do Banco de Dados', size=12, space_before=10)

add_body('O modelo foi projetado a partir do fluxo de trabalho científico do sistema, '
         'resultando em 9 tabelas:', space_after=4)

tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
tbl.autofit = True
add_table_row(tbl, ['Tabela', 'Responsabilidade'], bold=True, bg='D9E1F2')
tabelas = [
    ('usuarios',                  'Perfis de acesso com quatro categorias: gestor, colaborador, revisor, visitante'),
    ('especies_administrativo',   'Registro administrativo da espécie com progressão de status auditável'),
    ('especies_caracteristicas',  'Atributos morfológicos por parte da planta com campo de referência bibliográfica'),
    ('exemplares',                'Espécimes físicos de campo com código único, geolocalização e ciclo de aprovação'),
    ('especies_imagens',          'Fotografias das partes da planta vinculadas a um exemplar aprovado'),
    ('historico_alteracoes',      'Log imutável de todas as ações realizadas no sistema'),
    ('sugestoes_usuario',         'Canal para propostas de novas espécies ou correções'),
    ('tokens_recuperacao_senha',  'Tokens de uso único para redefinição de senha por e-mail'),
    ('tentativas_login',          'Proteção contra força bruta por IP'),
]
for t in tabelas:
    add_table_row(tbl, list(t))
doc.add_paragraph()

add_body('Relacionamentos principais:', bold=True, space_after=2)
rels = [
    'especies_caracteristicas → especies_administrativo (N:1, CASCADE DELETE)',
    'exemplares → especies_administrativo (N:1, CASCADE DELETE)',
    'exemplares → usuarios como especialista_id e cadastrado_por (N:1)',
    'especies_imagens → exemplares (N:1, SET NULL ao deletar exemplar)',
    'especies_imagens → usuarios como identificador, validador e coletor (N:1)',
    'historico_alteracoes → especies_administrativo e usuarios (N:1)',
]
for r in rels:
    add_bullet(r)

add_body('\nProgressão de status da espécie:', bold=True, space_after=2)
add_code([
    'sem_dados → dados_internet → descrita → registrada',
    '          → em_revisao → revisada → publicado',
    '                      ↘ contestado',
])
doc.add_paragraph()

add_body(
    'O código do exemplar é gerado automaticamente no formato XX000 (2 letras maiúsculas + '
    '3 dígitos sequenciais, ex.: KT001, BR042). Esse código é anotado na etiqueta de alumínio '
    'pregada no indivíduo físico no campo, garantindo o vínculo entre o espécime real e o '
    'registro digital.'
)

# ── 6.2 Implementação ─────────────────────────────────────────────────────────
add_heading('6.2 Implementação e Manipulação de Dados', size=12, space_before=10)

add_heading('DDL — Criação de tabelas', size=11, bold=True, space_before=6, space_after=3)
add_body('Tabela usuarios — controle de acesso por categoria com ENUM:', space_after=2)
add_code([
    'CREATE TABLE `usuarios` (',
    '  `id`                 INT(11)      NOT NULL AUTO_INCREMENT,',
    '  `nome`               VARCHAR(150) NOT NULL,',
    '  `email`              VARCHAR(150) NOT NULL,',
    '  `senha_hash`         VARCHAR(255) NOT NULL,',
    '  `categoria`          ENUM(\'gestor\',\'colaborador\',\'revisor\',\'visitante\')',
    '                       NOT NULL DEFAULT \'visitante\',',
    '  `status_verificacao` ENUM(\'pendente\',\'verificado\',\'bloqueado\')',
    '                       NOT NULL DEFAULT \'pendente\',',
    '  `ativo`              TINYINT(1)   NOT NULL DEFAULT 1,',
    '  `data_cadastro`      DATETIME     NOT NULL DEFAULT CURRENT_TIMESTAMP(),',
    '  PRIMARY KEY (`id`),',
    '  UNIQUE KEY `uq_email` (`email`)',
    ') ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;',
])
doc.add_paragraph()
add_body('Tabela exemplares — espécimes de campo com chaves estrangeiras:', space_after=2)
add_code([
    'CREATE TABLE `exemplares` (',
    '  `id`              INT(11)       NOT NULL AUTO_INCREMENT,',
    '  `codigo`          VARCHAR(6)    NOT NULL,',
    '  `especie_id`      INT(11)       NOT NULL,',
    '  `numero_etiqueta` VARCHAR(50)   DEFAULT NULL,',
    '  `latitude`        DECIMAL(10,8) DEFAULT NULL,',
    '  `longitude`       DECIMAL(11,8) DEFAULT NULL,',
    '  `cidade`          VARCHAR(150)  DEFAULT NULL,',
    '  `estado`          CHAR(2)       DEFAULT NULL,',
    '  `bioma`           VARCHAR(100)  DEFAULT NULL,',
    '  `especialista_id` INT(11)       NOT NULL,',
    '  `cadastrado_por`  INT(11)       NOT NULL,',
    '  `status`          ENUM(\'aguardando_revisao\',\'aprovado\',\'rejeitado\')',
    '                    NOT NULL DEFAULT \'aguardando_revisao\',',
    '  `data_cadastro`   DATETIME      NOT NULL DEFAULT CURRENT_TIMESTAMP(),',
    '  `data_revisao`    DATETIME      DEFAULT NULL,',
    '  `motivo_rejeicao` TEXT          DEFAULT NULL,',
    '  PRIMARY KEY (`id`),',
    '  UNIQUE KEY `uq_codigo` (`codigo`),',
    '  CONSTRAINT `fk_exemplar_especie`',
    '    FOREIGN KEY (`especie_id`) REFERENCES `especies_administrativo`(`id`)',
    '    ON DELETE CASCADE,',
    '  CONSTRAINT `fk_exemplar_especialista`',
    '    FOREIGN KEY (`especialista_id`) REFERENCES `usuarios`(`id`),',
    '  CONSTRAINT `fk_exemplar_cadastrador`',
    '    FOREIGN KEY (`cadastrado_por`)  REFERENCES `usuarios`(`id`)',
    ') ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;',
])
doc.add_paragraph()

add_heading('DML — Inserção de dados', size=11, bold=True, space_before=6, space_after=3)
add_body('Cadastro de espécies do Cerrado:', space_after=2)
add_code([
    'INSERT INTO `especies_administrativo` (nome_cientifico, status, prioridade)',
    'VALUES',
    "  ('Caryocar brasiliense',          'sem_dados', 'alta'),   -- pequizeiro",
    "  ('Handroanthus albus',            'sem_dados', 'alta'),   -- ipê-amarelo",
    "  ('Mauritia flexuosa',             'sem_dados', 'media'),  -- buriti",
    "  ('Stryphnodendron adstringens',   'sem_dados', 'media');  -- barbatimão",
])
doc.add_paragraph()
add_body('Registro de exemplar coletado em campo:', space_after=2)
add_code([
    'INSERT INTO `exemplares`',
    '  (codigo, especie_id, numero_etiqueta, latitude, longitude,',
    '   cidade, estado, bioma, descricao_local, especialista_id, cadastrado_por)',
    'VALUES',
    "  ('KT001', 1, 'ETQ-047',",
    '   -20.48290500, -54.61520000,',
    "   'Campo Grande', 'MS', 'Cerrado',",
    "   'Margem do córrego Segredo, árvore isolada em área de cerradão',",
    '   2, 3);',
])
doc.add_paragraph()

add_heading('DML — Atualização de dados', size=11, bold=True, space_before=6, space_after=3)
add_body('Aprovação de exemplar pelo especialista:', space_after=2)
add_code([
    'UPDATE `exemplares`',
    "SET   status       = 'aprovado',",
    '      data_revisao = NOW()',
    'WHERE id = 1',
    '  AND especialista_id = 2;',
])
doc.add_paragraph()
add_body('Avanço do status da espécie após descrição completa:', space_after=2)
add_code([
    'UPDATE `especies_administrativo`',
    "SET   status            = 'descrita',",
    '      data_descrita     = NOW(),',
    '      autor_descrita_id = 3',
    'WHERE id = 1',
    "  AND status = 'dados_internet';",
])
doc.add_paragraph()

add_heading('DML — Remoção de dados', size=11, bold=True, space_before=6, space_after=3)
add_body('Limpeza de tokens de senha expirados:', space_after=2)
add_code([
    'DELETE FROM `tokens_recuperacao_senha`',
    'WHERE expira_em < NOW()',
    '   OR usado = 1;',
])
doc.add_paragraph()
add_body('Limpeza de tentativas de login antigas (janela de 1 hora):', space_after=2)
add_code([
    'DELETE FROM `tentativas_login`',
    "WHERE criado_em < DATE_SUB(NOW(), INTERVAL 1 HOUR);",
])
doc.add_paragraph()

add_heading('DML — Consultas SQL', size=11, bold=True, space_before=6, space_after=3)
add_body('Total de imagens por espécie, agrupadas por status:', space_after=2)
add_code([
    'SELECT',
    '  ea.nome_cientifico,',
    '  ea.status,',
    '  COUNT(ei.id)                    AS total_imagens,',
    '  COUNT(DISTINCT ei.parte_planta) AS partes_distintas,',
    "  SUM(ei.status_validacao = 'aprovado') AS imagens_aprovadas",
    'FROM especies_administrativo ea',
    'LEFT JOIN especies_imagens ei ON ei.especie_id = ea.id',
    'GROUP BY ea.id, ea.nome_cientifico, ea.status',
    'ORDER BY partes_distintas DESC;',
])
doc.add_paragraph()
add_body('Exemplares aguardando revisão com dados do especialista:', space_after=2)
add_code([
    'SELECT',
    '  e.codigo,',
    '  ea.nome_cientifico,',
    '  e.cidade, e.estado, e.bioma,',
    '  e.data_cadastro,',
    '  u_esp.nome AS especialista,',
    '  u_col.nome AS coletado_por',
    'FROM exemplares e',
    'JOIN especies_administrativo ea ON ea.id = e.especie_id',
    'JOIN usuarios u_esp ON u_esp.id = e.especialista_id',
    'JOIN usuarios u_col ON u_col.id = e.cadastrado_por',
    "WHERE e.status = 'aguardando_revisao'",
    'ORDER BY e.data_cadastro ASC;',
])
doc.add_paragraph()
add_body('Histórico completo de alterações de uma espécie:', space_after=2)
add_code([
    'SELECT',
    '  ha.data_alteracao,',
    '  u.nome         AS responsavel,',
    '  u.categoria,',
    '  ha.tipo_acao,',
    '  ha.campo_alterado,',
    '  ha.valor_anterior,',
    '  ha.valor_novo',
    'FROM historico_alteracoes ha',
    'JOIN usuarios u ON u.id = ha.id_usuario',
    'WHERE ha.especie_id = 1',
    'ORDER BY ha.data_alteracao DESC;',
])
doc.add_paragraph()

# ── 6.3 Controle de Versão ────────────────────────────────────────────────────
add_heading('6.3 Uso do Controle de Versão', size=12, space_before=10)

add_body(
    'O projeto foi iniciado com git init e o primeiro commit registrado em fevereiro de 2026. '
    'O repositório acompanha toda a evolução do sistema, incluindo as alterações no esquema '
    'do banco de dados documentadas em arquivos SQL na pasta database/ e docs/sql/.'
)

add_heading('Convenções de mensagem de commit', size=11, bold=True, space_before=6, space_after=4)
tbl2 = doc.add_table(rows=1, cols=2)
tbl2.style = 'Table Grid'
add_table_row(tbl2, ['Prefixo', 'Uso'], bold=True, bg='D9E1F2')
convencoes = [
    ('feat:',     'Nova funcionalidade implementada'),
    ('fix:',      'Correção de bug'),
    ('refactor:', 'Reestruturação sem mudança de comportamento'),
    ('docs:',     'Atualização de documentação'),
    ('chore:',    'Manutenção (dependências, configurações)'),
    ('revert:',   'Reversão de commit anterior'),
]
for c in convencoes:
    add_table_row(tbl2, list(c))
doc.add_paragraph()

add_body('Exemplos de commits representativos do projeto:', bold=True, space_after=2)
add_code([
    '49fb6d5  initial commit: projeto penomato mvp',
    '8f64c7f  adiciona sistema completo de cadastro de características',
    '1d941e8  feat: sistema de cadastro botânico funcionando',
    '761833f  chore: remove perfil validador, adiciona extração de GPS',
    '8630055  refactor: simplifica seleção de foto (botão único)',
    '4dba957  fix: remove condição tentarExif indefinida',
    'a9e6dd8  revert: volta cadastrar_exemplar ao estado pré-mobile',
    '5048649  docs: registra diretrizes de dispositivo e evolução tecnológica',
    '09961ac  fix: corrige terminologia botânica nos atributos de folha',
    '',
    'Total de commits: 114',
])
doc.add_paragraph()

add_body('Link para o repositório:', bold=True, space_after=2)
add_body('(inserir URL do repositório GitHub após publicação)')

# ══════════════════════════════════════════════════════════════════════════════
# 7. RESULTADOS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('7. Resultados Preliminares', size=13, space_before=14)

add_body('O banco de dados penomato foi criado com sucesso no MariaDB 10.4 (XAMPP) e conta com:')
resultados = [
    '9 tabelas criadas e em operação',
    '117 espécies do Cerrado cadastradas na tabela especies_administrativo',
    '1 usuário gestor cadastrado (nilton.bakargi@ufms.br)',
    'Todas as restrições de integridade aplicadas (FKs, UNIQUEs, ENUMs, CHECK)',
    'Índices criados nos campos de busca mais frequentes (status, categoria, datas)',
]
for r in resultados:
    add_bullet(r)

add_body('\nCiclo de vida de uma espécie no banco:', bold=True, space_after=2)
add_code([
    '[Gestor cadastra]        → status: sem_dados',
    '[Colaborador descreve]   → status: dados_internet → descrita',
    '[Colaborador coleta]     → exemplar: aguardando_revisao',
    '[Especialista aprova]    → exemplar: aprovado',
    '[Fotos enviadas]         → status espécie: registrada',
    '[Artigo gerado]          → status: em_revisao',
    '[Especialista revisa]    → status: revisada → publicado',
])
doc.add_paragraph()
add_body(
    'Com 114 commits documentados, o histórico do Git demonstra a evolução progressiva do '
    'sistema: desde o commit inicial com a estrutura básica do banco, passando pela adição '
    'de tabelas de segurança, pela criação da tabela exemplares, até as correções de '
    'terminologia botânica nos atributos morfológicos.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 8. CONCLUSÃO
# ══════════════════════════════════════════════════════════════════════════════
add_heading('8. Conclusão', size=13, space_before=14)
add_body(
    'O desenvolvimento do banco de dados do Penomato demonstrou na prática os conceitos '
    'fundamentais de modelagem relacional, normalização e SQL aplicados a um problema real '
    'da área de ciências biológicas e engenharia florestal. O modelo resultante, com 9 tabelas '
    'e relacionamentos bem definidos, reflete fielmente o fluxo científico de catalogação: '
    'cada entidade corresponde a um ator ou artefato do processo real, e as restrições de '
    'integridade do banco garantem que regras de negócio críticas sejam aplicadas no nível '
    'do dado, não apenas na aplicação.'
)
add_body(
    'O controle de versão com Git provou ser indispensável para um projeto em evolução '
    'contínua. Os 114 commits documentados funcionam como um diário técnico do projeto, '
    'registrando não apenas o que foi alterado, mas por que foi alterado — informação '
    'essencial para qualquer desenvolvedor que precise compreender ou manter o sistema.'
)
add_body(
    'A combinação entre banco de dados relacional robusto e versionamento disciplinado forma '
    'a base técnica sobre a qual o Penomato poderá crescer: novas tabelas podem ser adicionadas '
    'como migrações numeradas, novas funcionalidades podem ser desenvolvidas em branches '
    'isoladas, e toda a equipe tem acesso a um histórico confiável e auditável do projeto.'
)

# ══════════════════════════════════════════════════════════════════════════════
# Salvar
# ══════════════════════════════════════════════════════════════════════════════
out = os.path.join(os.path.dirname(__file__), 'relatorio_modulo3.docx')
doc.save(out)
print(f'Arquivo gerado: {out}')
