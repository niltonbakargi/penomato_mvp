from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Margens ABNT
section = doc.sections[0]
section.page_width    = Cm(21)
section.page_height   = Cm(29.7)
section.top_margin    = Cm(3)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2)

# Estilos base
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = Pt(24)
normal.paragraph_format.space_after  = Pt(0)

def set_heading_style(style, size, bold=True):
    style.font.name = 'Times New Roman'
    style.font.size = Pt(size)
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(18)
    style.paragraph_format.space_after  = Pt(6)
    style.paragraph_format.line_spacing = Pt(24)

set_heading_style(doc.styles['Heading 1'], 14)
set_heading_style(doc.styles['Heading 2'], 12)
set_heading_style(doc.styles['Heading 3'], 12, bold=False)

def add_body(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def add_blank():
    p = doc.add_paragraph('')
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)

def add_h1(text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_h2(text):
    p = doc.add_heading(text, level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_h3(text):
    p = doc.add_heading(text, level=3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

def add_kv(label, text):
    p = doc.add_paragraph()
    r1 = p.add_run(label + ": ")
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r2 = p.add_run(text)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)

def add_ref(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(-1.25)
    p.paragraph_format.left_indent       = Cm(1.25)
    p.paragraph_format.line_spacing      = Pt(24)
    p.paragraph_format.space_after       = Pt(0)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

# ============================================================
# CAPA
# ============================================================
capa_lines = [
    ("UNIVERSIDADE ESTADUAL DE MATO GROSSO DO SUL", True, 12),
    ("CURSO DE ENGENHARIA FLORESTAL", True, 12),
    ("", False, 12),
    ("", False, 12),
    ("", False, 12),
    ("", False, 12),
    ("", False, 12),
    ("PENOMATO: PLATAFORMA COLABORATIVA PARA DOCUMENTACAO FITOMORFOLOGICA DE ESPECIES DO CERRADO", True, 14),
    ("", False, 12),
    ("", False, 12),
    ("Documento de Apresentacao do Prototipo -- Projeto de TCC em Tecnologia da Informacao (UFMS)", False, 12),
    ("Parceria: Departamento de Engenharia Florestal -- UEMS", False, 12),
    ("", False, 12),
    ("", False, 12),
    ("", False, 12),
    ("", False, 12),
    ("", False, 12),
    ("Autor: [Nome do Autor]", False, 12),
    ("Orientador: Prof. Norton Hayd Rego", False, 12),
    ("", False, 12),
    ("", False, 12),
    ("", False, 12),
    ("Campo Grande, MS -- 2026", False, 12),
]

for text, bold, size in capa_lines:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.bold = bold

doc.add_page_break()

# ============================================================
# RESUMO
# ============================================================
p_t = doc.add_paragraph("RESUMO")
p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in p_t.runs:
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
p_t.paragraph_format.space_before = Pt(0)
p_t.paragraph_format.space_after  = Pt(12)

resumo = (
    "O Penomato e uma plataforma web colaborativa destinada a coleta, organizacao e publicacao de dados "
    "fitomorfologicos de especies arboreas do Cerrado brasileiro. O sistema estrutura um fluxo cientifico "
    "completo -- do cadastro de especies a publicacao de fichas revisadas por especialistas -- integrando "
    "colaboradores de campo, professores orientadores e gestores institucionais. Desenvolvido como Trabalho "
    "de Conclusao de Curso em Tecnologia da Informacao pela UFMS, em parceria com o Departamento de "
    "Engenharia Florestal da UEMS, o projeto responde a escassez de dados digitais estruturados sobre a "
    "flora nativa do Cerrado e ao desaparecimento progressivo dos profissionais com conhecimento empirico "
    "de campo. O prototipo implementa cadastro morfologico assistido por Inteligencia Artificial, registro "
    "fotografico de exsicatas digitais vinculado a exemplares fisicos identificados por etiqueta, e fluxo "
    "de revisao por especialista com publicacao automatica. O sistema visa ser a base de dados necessaria "
    "para, em fases futuras, treinar modelos de visao computacional especializados em identificacao de "
    "especies nativas por bioma."
)
p_res = doc.add_paragraph(resumo)
p_res.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_res.paragraph_format.line_spacing = Pt(24)
p_res.paragraph_format.space_after  = Pt(12)
for r in p_res.runs:
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

p_kw = doc.add_paragraph()
p_kw.paragraph_format.line_spacing = Pt(24)
r1 = p_kw.add_run("Palavras-chave: ")
r1.bold = True
r1.font.name = 'Times New Roman'
r1.font.size = Pt(12)
r2 = p_kw.add_run("Cerrado; fitomorfologia; ciencia cidada; exsicata digital; identificacao de especies; Inteligencia Artificial; plataforma colaborativa.")
r2.font.name = 'Times New Roman'
r2.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# SUMARIO
# ============================================================
p_s = doc.add_paragraph("SUMARIO")
p_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in p_s.runs:
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
p_s.paragraph_format.space_after = Pt(18)

sumario_items = [
    "1  INTRODUCAO",
    "2  JUSTIFICATIVA",
    "3  OBJETIVOS",
    "4  METODOLOGIA E TECNOLOGIAS UTILIZADAS",
    "5  DESCRICAO DO SISTEMA",
    "   5.1  Perfis de usuario",
    "   5.2  Fluxo de status das especies",
    "   5.3  Cadastro morfologico assistido por IA",
    "   5.4  Modulo de exemplares de campo",
    "   5.5  Geracao e publicacao do artigo cientifico",
    "   5.6  Contestacao e reedicao",
    "6  DIFERENCIAIS DO PROJETO",
    "7  RESULTADOS ESPERADOS",
    "8  TRABALHOS FUTUROS",
    "9  REFERENCIAS",
]
for item in sumario_items:
    p = doc.add_paragraph(item)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# 1. INTRODUCAO
# ============================================================
add_h1("1  INTRODUCAO")

add_body(
    "A identificacao de especies arboreas em campo e um dos maiores desafios praticos da Engenharia Florestal. "
    "Profissionais conhecidos como mateiros -- detentores de conhecimento empirico profundo sobre a flora nativa -- "
    "estao progressivamente desaparecendo, e os projetos de inventario florestal enfrentam crescente dificuldade "
    "para encontrar mao de obra capacitada. No Cerrado, bioma com a segunda maior biodiversidade do Brasil e elevado "
    "grau de endemismo, esse problema e particularmente critico: a vegetacao e diversificada, os padroes morfologicos "
    "variam sazonalmente e os recursos de identificacao digital disponiveis sao escassos ou fragmentados."
)
add_blank()
add_body(
    "O projeto Penomato nasceu dessa lacuna. A ideia original surgiu de um engenheiro florestal que, ao realizar "
    "inventarios em campo, reconheceu a inexistencia de uma ferramenta digital capaz de identificar especies "
    "pelo conjunto de caracteristicas morfologicas observaveis no individuo. As solucoes tentadas -- desde planilhas "
    "Excel ate chaves dicotomicas digitalizadas -- esbarravam sempre no mesmo obstaculo: a ausencia de dados "
    "estruturados e validados sobre as especies do Cerrado em formato adequado para computacao."
)
add_blank()
add_body(
    "A evolucao da solucao levou a uma conclusao fundamental: antes de construir um identificador, era necessario "
    "construir a infraestrutura capaz de gerar os dados que tal identificador precisaria. O Penomato e, portanto, "
    "essa infraestrutura -- uma plataforma colaborativa que organiza o processo cientifico de documentacao "
    "fitomorfologica, da coleta em campo a publicacao revisada, criando progressivamente uma base de dados "
    "robusta e rastreaval."
)
add_blank()
add_body(
    "Este documento descreve o prototipo funcional desenvolvido como Trabalho de Conclusao de Curso (TCC) em "
    "Tecnologia da Informacao pela Universidade Federal de Mato Grosso do Sul (UFMS), em parceria com o "
    "Departamento de Engenharia Florestal da Universidade Estadual de Mato Grosso do Sul (UEMS). O prototipo "
    "implementa o fluxo completo de documentacao de especies, desde o cadastro morfologico ate a publicacao "
    "de fichas cientificas revisadas por especialistas."
)

# ============================================================
# 2. JUSTIFICATIVA
# ============================================================
add_h1("2  JUSTIFICATIVA")

add_body(
    "O Cerrado abriga aproximadamente 11.000 especies de plantas vasculares, das quais cerca de 44% sao endemicas "
    "(MYERS et al., 2000). Apesar de sua importancia ecologica e das pressoes crescentes de desmatamento, "
    "a documentacao digital estruturada desse acervo biologico permanece fragmentada. As ferramentas existentes "
    "-- como a Flora do Brasil (JBRJ/REFLORA) -- oferecem listagens taxonomicas mas nao fornecem dados morfologicos "
    "estruturados e fotografias de campo com rastreabilidade do individuo fisico, que sao os elementos essenciais "
    "para treinamento de modelos de visao computacional."
)
add_blank()
add_body(
    "A pesquisa com profissionais da area revelou que o principal motivador para contribuicao voluntaria e o "
    "reconhecimento academico. Engenheiros florestais, biologos e estudantes de graduacao estao dispostos a "
    "contribuir com dados de qualidade se receberem credito cientifico pelo trabalho realizado. O Penomato "
    "incorpora esse incentivo diretamente na arquitetura do sistema: cada especie documentada gera um artigo "
    "cientifico com creditacao de todos os contribuidores -- o coletor de dados, o fotografo de campo e o "
    "especialista revisor."
)
add_blank()
add_body(
    "Adicionalmente, o sistema foi projetado para minimizar a friccao de entrada de dados. A geracao atual de "
    "colaboradores -- estudantes universitarios e profissionais jovens -- apresenta baixa tolerancia a formularios "
    "extensos e processos manuais repetitivos. A integracao de Inteligencia Artificial para preenchimento "
    "automatico do formulario morfologico, com o colaborador assumindo o papel de curador em vez de digitador, "
    "e uma decisao de produto deliberada para maximizar o engajamento e a retencao na plataforma."
)
add_blank()
add_body(
    "Do ponto de vista institucional, a parceria entre UFMS e UEMS posiciona o projeto em um contexto de "
    "producao cientifica real. O uso do Penomato por alunos do curso de Engenharia Florestal da UEMS como "
    "ferramenta didatica de campo -- com os professores assumindo o papel de especialistas revisores -- representa "
    "um modelo de adocao imediato que valida o sistema com dados reais antes mesmo da defesa do TCC."
)

# ============================================================
# 3. OBJETIVOS
# ============================================================
add_h1("3  OBJETIVOS")

add_h2("3.1  Objetivo Geral")
add_body(
    "Desenvolver uma plataforma web colaborativa para coleta, organizacao, validacao e publicacao de dados "
    "fitomorfologicos de especies arboreas do Cerrado, estruturando um fluxo cientifico rastreaval do individuo "
    "fisico de campo ao artigo publicado."
)

add_h2("3.2  Objetivos Especificos")
add_blank()
objetivos = [
    "Implementar cadastro morfologico assistido por Inteligencia Artificial com revisao humana obrigatoria;",
    "Estruturar o registro fotografico de exsicatas digitais vinculado a exemplares fisicos identificados por etiqueta numerada;",
    "Desenvolver fluxo de revisao por especialista com aprovacao, rejeicao motivada e publicacao automatica;",
    "Criar sistema de contestacao e reedicao de artigos cientificos com controle de versoes;",
    "Disponibilizar ficha publica por especie com galeria fotografica e creditacao de todos os contribuidores;",
    "Estabelecer a base de dados estruturada necessaria para, em etapas futuras, treinar modelos de visao computacional especializados em identificacao de especies nativas.",
]
for obj in objetivos:
    add_bullet(obj)

# ============================================================
# 4. METODOLOGIA
# ============================================================
add_h1("4  METODOLOGIA E TECNOLOGIAS UTILIZADAS")

add_body(
    "O desenvolvimento seguiu metodologia agil com entregas incrementais, priorizando o fluxo cientifico "
    "completo de ponta a ponta antes da adicao de funcionalidades auxiliares. As decisoes de arquitetura foram "
    "orientadas por tres principios: rastreabilidade cientifica (cada dado tem origem, autor e data registrados), "
    "qualidade garantida por processo (nenhum dado chega a publicacao sem revisao de especialista) e "
    "baixo custo de adocao (interface web responsiva, sem necessidade de instalacao de aplicativo)."
)
add_blank()
add_body("A stack tecnologica adotada foi:")

tecnologias = [
    "Back-end: PHP 8 com arquitetura MVC (Model-View-Controller), MySQL via XAMPP;",
    "Front-end: HTML5, CSS3 com Bootstrap 5, JavaScript -- interface responsiva, compativel com dispositivos moveis de campo;",
    "Mapas interativos: biblioteca Leaflet.js com tiles OpenStreetMap, para georreferenciamento de exemplares;",
    "Inteligencia Artificial: API Claude (Anthropic) para preenchimento automatico de atributos morfologicos, com suporte a OpenAI, Gemini e DeepSeek como alternativas;",
    "Infraestrutura: servidor Apache local (XAMPP) para desenvolvimento; estrutura pronta para migracao para servidor de producao.",
]
for tec in tecnologias:
    add_bullet(tec)

add_blank()
add_body(
    "O banco de dados foi modelado para suportar tanto o fluxo atual quanto as expansoes planejadas. "
    "Todos os eventos relevantes do sistema -- alteracoes de status, aprovacoes, rejeicoes, confirmacoes "
    "de atributos -- sao registrados em tabela de auditoria com usuario, data e conteudo anterior e posterior, "
    "garantindo rastreabilidade cientifica completa."
)

# ============================================================
# 5. DESCRICAO DO SISTEMA
# ============================================================
add_h1("5  DESCRICAO DO SISTEMA")

add_body(
    "O Penomato organiza seu fluxo em torno de tres entidades centrais: a especie (objeto de documentacao), "
    "o exemplar (o individuo fisico de campo) e o artigo (a publicacao resultante). O sistema gerencia "
    "a progressao dessas entidades por um conjunto de status bem definidos, com transicoes controladas "
    "e notificacoes automaticas para os usuarios responsaveis."
)

add_h2("5.1  Perfis de Usuario")
add_blank()

perfis = [
    ("Gestor", "Administra o catalogo de especies de interesse, atribui colaboradores e dispensa partes nao disponiveis para coleta. Tem visao completa do estado de todas as especies no sistema."),
    ("Colaborador", "Insere dados morfologicos a partir de fontes da internet, confirma atributos por revisao individual, cadastra exemplares de campo e envia fotografias das partes da planta."),
    ("Especialista (Revisor)", "Revisa e aprova exemplares de campo antes que as fotos possam ser enviadas; revisa os artigos gerados antes da publicacao; pode rejeitar com motivo fundamentado."),
    ("Visitante (Publico)", "Acessa fichas publicas das especies publicadas, sem necessidade de cadastro."),
]
for nome, descricao in perfis:
    add_kv(nome, descricao)

add_h2("5.2  Fluxo de Status das Especies")

add_body(
    "Cada especie percorre uma progressao de status que reflete seu estado de documentacao. Os status sao:"
)
add_blank()

status_list = [
    ("Sem Dados", "especie cadastrada pelo gestor, aguardando contribuicao;"),
    ("Dados da Internet", "colaborador inseriu descricao morfologica e imagens de referencia;"),
    ("Identificada", "todos os atributos morfologicos foram confirmados individualmente por um colaborador;"),
    ("Registrada", "todas as partes da planta foram fotografadas no campo (ou formalmente dispensadas);"),
    ("Aguardando Revisao", "especie esta identificada e registrada; artigo gerado e na fila do especialista;"),
    ("Revisada / Publicada", "especialista aprovou o artigo; ficha publica disponivel;"),
    ("Em Contestacao", "identificacao questionada apos publicacao; em processo de correcao."),
]
for status, desc in status_list:
    p = doc.add_paragraph()
    r1 = p.add_run(status + ": ")
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r2 = p.add_run(desc)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent  = Cm(1.25)
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)

add_blank()
add_body(
    "Os caminhos de Identificacao (confirmacao de atributos) e de Registro (fotografias de campo) correm "
    "em paralelo e de forma independente. A geracao do artigo so e desbloqueada quando ambos os caminhos "
    "estiverem completos para a mesma especie. O sistema emite notificacoes informando o que falta em cada caso."
)

add_h2("5.3  Cadastro Morfologico Assistido por Inteligencia Artificial")

add_body(
    "O colaborador acessa o formulario de cadastro morfologico e aciona o botao Consultar IA. "
    "O sistema envia uma requisicao a API de linguagem (Claude/Anthropic por padrao), que retorna "
    "uma descricao estruturada com os atributos morfologicos da especie: folha, flor, fruto, caule, "
    "semente, habito, distribuicao geografica, sinonimias e referencias bibliograficas. O colaborador "
    "revisa cada campo individualmente antes de confirmar o envio, assumindo papel de curador em vez "
    "de digitador."
)
add_blank()
add_body(
    "Em etapa posterior, o colaborador acessa a tela de confirmacao de caracteristicas e revisa "
    "cada atributo registrado um a um, podendo confirma-lo como correto ou substitui-lo pelo valor "
    "verificado. Somente quando todos os atributos estiverem confirmados ou corrigidos a especie "
    "avanca para o status Identificada. Esse processo garante que nenhum dado gerado automaticamente "
    "chega ao artigo sem validacao humana explicita."
)

add_h2("5.4  Modulo de Exemplares de Campo")

add_body(
    "O conceito de exemplar e central para a rastreabilidade cientifica do sistema. Um exemplar "
    "corresponde a um individuo fisico especifico -- uma planta identificada no campo por uma etiqueta "
    "de aluminio numerada pregada em seu tronco. Todas as fotografias das partes (folha, flor, fruto, "
    "caule, semente, habito) devem ser vinculadas ao mesmo exemplar, garantindo que nunca se misturem "
    "partes de individuos diferentes em uma mesma colecao."
)
add_blank()
add_body(
    "O fluxo do exemplar e o seguinte: o colaborador cadastra o exemplar com geolocalizacao "
    "(capturada pelo navegador ou inserida manualmente), foto de identificacao geral da planta e "
    "selecao do especialista orientador. O sistema gera um codigo unico no formato XX000 "
    "(duas letras aleatorias mais sequencial numerico, ex.: KT001). O exemplar entra com status "
    "Aguardando Revisao e os uploads de partes ficam bloqueados ate aprovacao do especialista."
)
add_blank()
add_body(
    "O especialista acessa seu painel de revisao, visualiza a foto de identificacao e as "
    "informacoes de localizacao em mapa interativo, e aprova ou rejeita com motivo fundamentado. "
    "Uma vez aprovado o exemplar, os colaboradores podem enviar as fotos das partes vinculadas a ele. "
    "Multiplos colaboradores podem contribuir com partes diferentes do mesmo exemplar em momentos "
    "distintos, permitindo coleta incremental em campo."
)

add_h2("5.5  Geracao e Publicacao do Artigo Cientifico")

add_body(
    "Quando uma especie atinge simultaneamente os status Identificada e Registrada, o sistema "
    "habilita a geracao do artigo. O colaborador responsavel (ou o gestor) aciona explicitamente "
    "essa geracao apos confirmacao. O artigo produzido contem: nome cientifico e popular, familia "
    "taxonomica, sinonimias, descricao morfologica completa por parte da planta, galeria fotografica "
    "das exsicatas com metadados (coletor, data, codigo do exemplar, bioma), referencias bibliograficas "
    "e lista completa de todos os contribuidores."
)
add_blank()
add_body(
    "O artigo gerado e encaminhado ao especialista orientador do exemplar, que acessa o painel de "
    "revisao com ferramentas de visualizacao de imagens (zoom, ajuste de brilho, contraste e saturacao -- "
    "apenas para analise, sem alteracao dos arquivos). O especialista pode aprovar com parecer "
    "registrado ou rejeitar com motivo, retornando a especie ao status correspondente ao problema "
    "identificado. Apos aprovacao, a publicacao e automatica: a ficha publica da especie e gerada "
    "e disponibilizada sem necessidade de acao adicional."
)

add_h2("5.6  Contestacao e Reedicao")

add_body(
    "Apos a publicacao, qualquer colaborador ou especialista pode abrir uma contestacao, registrando "
    "o motivo. Se a contestacao for aceita, a identificacao e corrigida enquanto as fotografias "
    "permanecem associadas ao exemplar fisico (que nao muda -- a planta ainda existe no campo). "
    "O artigo e regenerado com a identificacao corrigida, passa por nova revisao do especialista "
    "e e republicado como nova edicao, com historico de versoes mantido."
)

# ============================================================
# 6. DIFERENCIAIS
# ============================================================
add_h1("6  DIFERENCIAIS DO PROJETO")
add_blank()

diferenciais = [
    "Rastreabilidade completa individuo-parte-artigo: cada fotografia esta ligada a um exemplar fisico especifico, identificado por etiqueta e geolocalizado, com registro do coletor e da data de coleta;",
    "IA como auxiliar, nao como substituto: a Inteligencia Artificial acelera o preenchimento morfologico, mas cada atributo gerado automaticamente precisa de confirmacao humana explicita antes de avancar no fluxo;",
    "Fluxo de revisao estruturado: nenhuma especie e publicada sem aprovacao de especialista, garantindo qualidade cientifica dos dados;",
    "Dados rotulados por parte e por bioma: estrutura de armazenamento das imagens ja esta preparada para uso futuro em treinamento de modelos de visao computacional;",
    "Incentivo a colaboracao por reconhecimento academico: cada contribuidor aparece com credito explicito no artigo publicado -- o sistema produz dados cientificos com autoria rastreaval;",
    "Flora do Cerrado integrada: o modulo publico traz consulta a base REFLORA/JBRJ com filtros por bioma, complementando as fichas produzidas pelos colaboradores.",
]
for d in diferenciais:
    add_bullet(d)

# ============================================================
# 7. RESULTADOS ESPERADOS
# ============================================================
add_h1("7  RESULTADOS ESPERADOS")

add_body(
    "A adocao do Penomato pelo curso de Engenharia Florestal da UEMS como ferramenta didatica de campo "
    "produzira, como resultado direto, um banco de dados de especies do Cerrado com descricoes morfologicas "
    "validadas e fotografias de exsicatas digitais georreferenciadas, vinculadas a individuos fisicos "
    "identificaveis. Cada especie cadastrada de ponta a ponta gerara um artigo cientifico disponivel "
    "publicamente, com creditacao dos alunos e professores envolvidos."
)
add_blank()
add_body(
    "Do ponto de vista tecnologico, espera-se validar o fluxo completo do sistema com dados reais -- "
    "identificando gargalos de usabilidade e ajustando o processo para o contexto real de uso em campo. "
    "O prototipo funcionando com dados reais e o argumento principal para a proxima fase de expansao "
    "do projeto para outros biomas e outras instituicoes."
)
add_blank()
add_body(
    "Para a UFMS, o resultado e a defesa de um TCC com contribuicao real ao problema da documentacao "
    "da flora nativa brasileira, com software funcional e base de dados gerada em colaboracao institucional. "
    "Para a UEMS e o curso de Engenharia Florestal, o resultado e uma ferramenta didatica original que "
    "integra atividade de campo, producao cientifica e tecnologia de forma inedita na regiao."
)

# ============================================================
# 8. TRABALHOS FUTUROS
# ============================================================
add_h1("8  TRABALHOS FUTUROS")

add_body(
    "O Penomato foi projetado como fundacao de dados para uma visao de longo prazo mais ampla. "
    "As etapas planejadas para alem do MVP sao:"
)
add_blank()

futuros = [
    ("Modelo de visao computacional especializado por bioma",
     "Com volume suficiente de imagens rotuladas por parte da planta (folha, caule, flor, fruto, semente) e por bioma, "
     "sera possivel treinar um modelo de identificacao de especies baseado em imagens. A inovacao metodologica proposta "
     "nao e a identificacao por uma unica foto, mas a analise combinada de multiplas partes do mesmo individuo, "
     "comparada com especies do bioma local -- reduzindo drasticamente os falsos positivos caracteristicos de "
     "identificadores de flora genericos."),
    ("Aplicativo movel com modo offline",
     "Interface nativa para dispositivos moveis com funcionamento offline em campo, sincronizacao automatica "
     "ao retornar a area com conectividade, e captura GPS nativa integrada ao cadastro do exemplar."),
    ("Sistema de gamificacao",
     "Contador de especies cadastradas por colaborador, ranking de colaboradores mais ativos, notificacoes "
     "de publicacao de especies contribuidas e medalhas por meta atingida -- mecanismos de engajamento "
     "inspirados em plataformas de ciencia cidada como o iNaturalist."),
    ("Expansao para todos os biomas brasileiros",
     "Apos validacao no Cerrado, expansao do modelo para Amazonia, Mata Atlantica, Pantanal, Caatinga e Pampa, "
     "em parceria com universidades e institutos de pesquisa de cada regiao."),
    ("Plataforma de educacao ambiental",
     "Convergencia de todos os modulos em uma plataforma de educacao ambiental para diferentes publicos: "
     "do pesquisador ao estudante do ensino fundamental, com fichas ilustradas, mapas interativos de "
     "ocorrencia e, futuramente, um aplicativo de caca a especies nativas em campo utilizando o modelo "
     "de visao computacional treinado com os dados da plataforma."),
]

for titulo, texto in futuros:
    add_h3(titulo)
    add_body(texto)

# ============================================================
# 9. REFERENCIAS
# ============================================================
add_h1("9  REFERENCIAS")

refs = [
    "MYERS, N. et al. Biodiversity hotspots for conservation priorities. Nature, v. 403, p. 853-858, 2000.",
    "JBRJ -- JARDIM BOTANICO DO RIO DE JANEIRO. Flora e Funga do Brasil. Disponivel em: <floradobrasil.jbrj.gov.br>. Acesso em: mai. 2026.",
    "REFLORA -- Rede Brasileira de Herbarios. Sistema de Informacao sobre a Biodiversidade Brasileira -- SiBBr. Disponivel em: <reflora.jbrj.gov.br>. Acesso em: mai. 2026.",
    "INATURALIST. Plataforma de Ciencia Cidada para Registro da Biodiversidade. Disponivel em: <www.inaturalist.org>. Acesso em: mai. 2026.",
    "ANTHROPIC. Claude API Documentation. Disponivel em: <docs.anthropic.com>. Acesso em: mai. 2026.",
    "BOOTSTRAP. Bootstrap 5 -- The most popular HTML, CSS, and JS library. Disponivel em: <getbootstrap.com>. Acesso em: mai. 2026.",
    "LEAFLET. Leaflet -- an open-source JavaScript library for interactive maps. Disponivel em: <leafletjs.com>. Acesso em: mai. 2026.",
]

for ref in refs:
    add_ref(ref)

# Salvar
output_path = r'C:\xampp\htdocs\penomato_mvp\docs\penomato_descricao_projeto.docx'
doc.save(output_path)
print("Salvo em:", output_path)
